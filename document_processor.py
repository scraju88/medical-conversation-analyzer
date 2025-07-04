import os
import PyPDF2
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document as LangchainDocument
import chromadb
from typing import List, Dict, Any
import hashlib
import json
import openai

class DocumentProcessor:
    def __init__(self, knowledge_base_path: str = "knowledge_base"):
        self.knowledge_base_path = knowledge_base_path
        self.embeddings = OpenAIEmbeddings()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
        # Create knowledge base directory
        if not os.path.exists(knowledge_base_path):
            os.makedirs(knowledge_base_path)
        
        # Initialize ChromaDB
        self.vectorstore = Chroma(
            persist_directory=knowledge_base_path,
            embedding_function=self.embeddings
        )
        
        # Update OpenAI client
        client = openai.OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Error extracting text from TXT: {e}")
            return ""
    
    def process_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Process a document and add it to the knowledge base"""
        try:
            # Extract text based on file type
            if file_type.lower() == 'pdf':
                text = self.extract_text_from_pdf(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                text = self.extract_text_from_docx(file_path)
            elif file_type.lower() == 'txt':
                text = self.extract_text_from_txt(file_path)
            else:
                return {"success": False, "error": f"Unsupported file type: {file_type}"}
            
            if not text.strip():
                return {"success": False, "error": "No text extracted from document"}
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Create documents for vector store
            documents = []
            for i, chunk in enumerate(chunks):
                doc = LangchainDocument(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "chunk_id": i,
                        "file_type": file_type,
                        "total_chunks": len(chunks)
                    }
                )
                documents.append(doc)
            
            # Add to vector store
            self.vectorstore.add_documents(documents)
            
            # Save document metadata
            self._save_document_metadata(file_path, file_type, len(chunks))
            
            return {
                "success": True,
                "chunks_created": len(chunks),
                "file_path": file_path,
                "file_type": file_type
            }
            
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _save_document_metadata(self, file_path: str, file_type: str, chunk_count: int):
        """Save metadata about processed documents"""
        metadata_file = os.path.join(self.knowledge_base_path, "documents_metadata.json")
        
        # Load existing metadata
        if os.path.exists(metadata_file):
            with open(metadata_file, 'r') as f:
                metadata = json.load(f)
        else:
            metadata = []
        
        # Add new document metadata
        doc_metadata = {
            "file_path": file_path,
            "file_type": file_type,
            "chunk_count": chunk_count,
            "processed_at": str(os.path.getmtime(file_path))
        }
        
        metadata.append(doc_metadata)
        
        # Save updated metadata
        with open(metadata_file, 'w') as f:
            json.dump(metadata, f, indent=2)
    
    def search_knowledge_base(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """Search the knowledge base for relevant information"""
        try:
            # Search the vector store
            docs = self.vectorstore.similarity_search(query, k=top_k)
            
            results = []
            for doc in docs:
                results.append({
                    "content": doc.page_content,
                    "metadata": doc.metadata,
                    "score": 0.8  # Placeholder score
                })
            
            return results
        except Exception as e:
            print(f"Error searching knowledge base: {e}")
            return []
    
    def get_knowledge_base_stats(self) -> Dict[str, Any]:
        """Get statistics about the knowledge base"""
        try:
            metadata_file = os.path.join(self.knowledge_base_path, "documents_metadata.json")
            
            if os.path.exists(metadata_file):
                with open(metadata_file, 'r') as f:
                    metadata = json.load(f)
                
                total_documents = len(metadata)
                total_chunks = sum(doc.get("chunk_count", 0) for doc in metadata)
                
                # Count by file type
                file_types = {}
                for doc in metadata:
                    file_type = doc.get("file_type", "unknown")
                    file_types[file_type] = file_types.get(file_type, 0) + 1
                
                return {
                    "total_documents": total_documents,
                    "total_chunks": total_chunks,
                    "file_types": file_types
                }
            else:
                return {
                    "total_documents": 0,
                    "total_chunks": 0,
                    "file_types": {}
                }
        except Exception as e:
            print(f"Error getting knowledge base stats: {e}")
            return {
                "total_documents": 0,
                "total_chunks": 0,
                "file_types": {}
            }
    
    def clear_knowledge_base(self):
        """Clear the entire knowledge base"""
        try:
            # Clear vector store
            self.vectorstore.delete_collection()
            
            # Remove metadata file
            metadata_file = os.path.join(self.knowledge_base_path, "documents_metadata.json")
            if os.path.exists(metadata_file):
                os.remove(metadata_file)
            
            return {"success": True, "message": "Knowledge base cleared successfully"}
        except Exception as e:
            return {"success": False, "error": str(e)}

def analyze_with_openai(transcription):
    """Analyze transcription using OpenAI GPT-4 with RAG context"""
    try:
        # Get relevant context from knowledge base
        context = get_relevant_context(transcription)
        
        prompt = f"""
        ... (your prompt as before) ...
        """
        
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a medical professional assistant. Provide accurate, professional medical analysis using available knowledge base information."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.3
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error analyzing with OpenAI: {str(e)}" 