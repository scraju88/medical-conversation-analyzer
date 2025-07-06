from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
import os
import json
import datetime
import speech_recognition as sr
from pydub import AudioSegment
import openai
import google.generativeai as genai
from dotenv import load_dotenv
import tempfile
import uuid
from document_processor import DocumentProcessor
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.messages import SystemMessage, HumanMessage
from langchain_chroma import Chroma
from langchain_core.documents import Document as LangchainDocument

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)

# Database configuration
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///medical_conversations.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

# Configure API keys
openai.api_key = os.getenv('OPENAI_API_KEY')
genai.configure(api_key=os.getenv('GEMINI_API_KEY'))

# Initialize document processor
document_processor = DocumentProcessor()

# Initialize TemplateVectorStore
class TemplateVectorStore:
    def __init__(self, persist_directory: str = "template_vectors"):
        self.persist_directory = persist_directory
        self.embeddings = OpenAIEmbeddings()
        self.vectorstore = Chroma(
            persist_directory=persist_directory,
            embedding_function=self.embeddings
        )
    def add_template(self, template_id: str, content: str, metadata: dict):
        doc = LangchainDocument(
            page_content=content,
            metadata={"template_id": template_id, **metadata}
        )
        self.vectorstore.add_documents([doc])
    def search_templates(self, query: str, top_k: int = 5):
        docs = self.vectorstore.similarity_search(query, k=top_k)
        return docs

template_vectorstore = TemplateVectorStore()

# Database Models
class MedicalConversation(db.Model):
    id = db.Column(db.String(36), primary_key=True)
    patient_name = db.Column(db.String(100))
    physician_name = db.Column(db.String(100))
    conversation_date = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    audio_file_path = db.Column(db.String(500))
    transcription = db.Column(db.Text)
    history_of_present_illness = db.Column(db.Text)
    assessment_plan = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)

class Template(db.Model):
    id = db.Column(db.String(36), primary_key=True)
    name = db.Column(db.String(200), nullable=False)
    content = db.Column(db.Text, nullable=False)
    template_metadata = db.Column(db.Text)  # JSON string for tags, specialty, etc.
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)

# Create database tables
with app.app_context():
    db.create_all()

def transcribe_audio(audio_file_path):
    """Transcribe audio file to text using Google Speech Recognition"""
    recognizer = sr.Recognizer()
    
    # Convert audio to WAV if needed
    audio = AudioSegment.from_file(audio_file_path)
    wav_path = audio_file_path.replace('.m4a', '.wav').replace('.mp3', '.wav')
    audio.export(wav_path, format="wav")
    
    with sr.AudioFile(wav_path) as source:
        audio_data = recognizer.record(source)
        try:
            text = recognizer.recognize_google(audio_data)
            return text
        except sr.UnknownValueError:
            return "Could not understand audio"
        except sr.RequestError as e:
            return f"Error with speech recognition service: {e}"
        finally:
            # Clean up temporary WAV file
            if os.path.exists(wav_path):
                os.remove(wav_path)

def get_relevant_context(transcription):
    """Get relevant context from knowledge base"""
    try:
        # Search for relevant information based on transcription
        relevant_docs = document_processor.search_knowledge_base(transcription, top_k=3)
        
        if relevant_docs:
            context = "\n\nRELEVANT KNOWLEDGE BASE INFORMATION:\n"
            for i, doc in enumerate(relevant_docs, 1):
                context += f"{i}. {doc['content']}\n"
            return context
        else:
            return ""
    except Exception as e:
        print(f"Error getting context: {e}")
        return ""

def analyze_with_openai(transcription):
    """Analyze transcription using OpenAI GPT-4 with RAG context via LangChain"""
    try:
        context = get_relevant_context(transcription)
        prompt = f"""
        You are a medical professional analyzing a conversation between a patient and physician. 
        Please provide a structured medical report with the following sections:

        CONVERSATION TRANSCRIPT:
        {transcription}

        {context}

        Please analyze this conversation and provide a structured response with clear section headers:

        1. HISTORY OF PRESENT ILLNESS (HPI):
        - Chief complaint
        - History of present illness
        - Relevant symptoms and their timeline
        - Pertinent positive and negative findings

        2. ASSESSMENT AND PLAN:
        - Differential diagnosis
        - Assessment of current condition
        - Treatment plan
        - Follow-up recommendations
        - Any additional tests or referrals needed

        Use the knowledge base information to enhance your analysis with relevant medical guidelines, 
        protocols, or best practices that apply to this case.

        IMPORTANT: Format your response exactly as follows:
        Start with \"1. HISTORY OF PRESENT ILLNESS (HPI):\" on its own line
        Then provide the HPI content
        Then start a new section with \"2. ASSESSMENT AND PLAN:\" on its own line
        Then provide the assessment and plan content
        """
        llm = ChatOpenAI(
            model="gpt-4",
            openai_api_key=os.getenv("OPENAI_API_KEY"),
            temperature=0.3,
            max_tokens=2000,
        )
        messages = [
            SystemMessage(content="You are a medical professional assistant. Provide accurate, professional medical analysis using available knowledge base information."),
            HumanMessage(content=prompt)
        ]
        response = llm.invoke(messages)
        return response.content
    except Exception as e:
        return f"Error analyzing with OpenAI: {str(e)}"

def analyze_with_gemini(transcription):
    """Analyze transcription using Google Gemini with RAG context"""
    try:
        model = genai.GenerativeModel('gemini-pro')
        
        # Get relevant context from knowledge base
        context = get_relevant_context(transcription)
        
        prompt = f"""
        You are a medical professional analyzing a conversation between a patient and physician. 
        Please provide a structured medical report with the following sections:

        CONVERSATION TRANSCRIPT:
        {transcription}

        {context}

        Please analyze this conversation and provide a structured response with clear section headers:

        1. HISTORY OF PRESENT ILLNESS (HPI):
        - Chief complaint
        - History of present illness
        - Relevant symptoms and their timeline
        - Pertinent positive and negative findings

        2. ASSESSMENT AND PLAN:
        - Differential diagnosis
        - Assessment of current condition
        - Treatment plan
        - Follow-up recommendations
        - Any additional tests or referrals needed

        Use the knowledge base information to enhance your analysis with relevant medical guidelines, 
        protocols, or best practices that apply to this case.

        IMPORTANT: Format your response exactly as follows:
        Start with "1. HISTORY OF PRESENT ILLNESS (HPI):" on its own line
        Then provide the HPI content
        Then start a new section with "2. ASSESSMENT AND PLAN:" on its own line
        Then provide the assessment and plan content
        """
        
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error analyzing with Gemini: {str(e)}"

def fill_template_with_llm(template_content, hpi, assessment_plan):
    prompt = f"""
    You are a medical documentation assistant. Here is a template for a medical note:

    TEMPLATE:
    {template_content}

    Here is the patient's History of Present Illness (HPI):
    {hpi}

    Here is the Assessment and Plan:
    {assessment_plan}

    Please fill out the template above with the provided HPI and Assessment & Plan in the most appropriate locations, using your best judgment. Return the completed note.
    """
    llm = ChatOpenAI(
        model="gpt-4",
        openai_api_key=os.getenv("OPENAI_API_KEY"),
        temperature=0.3,
        max_tokens=2000,
    )
    messages = [
        SystemMessage(content="You are a helpful assistant for medical documentation."),
        HumanMessage(content=prompt)
    ]
    response = llm.invoke(messages)
    return response.content

@app.route('/')
def index():
    """Serve the web interface"""
    return render_template('index.html')

@app.route('/api/upload', methods=['POST'])
def upload_audio():
    """Upload and process audio file"""
    try:
        if 'audio' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400
        
        audio_file = request.files['audio']
        patient_name = request.form.get('patient_name', 'Unknown Patient')
        physician_name = request.form.get('physician_name', 'Unknown Physician')
        llm_provider = request.form.get('llm_provider', 'openai')  # openai or gemini
        
        if audio_file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Generate unique ID for the conversation
        conversation_id = str(uuid.uuid4())
        
        # Save audio file
        upload_folder = 'uploads'
        if not os.path.exists(upload_folder):
            os.makedirs(upload_folder)
        
        audio_path = os.path.join(upload_folder, f"{conversation_id}_{audio_file.filename}")
        audio_file.save(audio_path)
        
        # Transcribe audio
        transcription = transcribe_audio(audio_path)
        
        # Analyze with LLM
        if llm_provider == 'gemini':
            analysis = analyze_with_gemini(transcription)
        else:
            analysis = analyze_with_openai(transcription)
        
        print("AI RAW RESPONSE:")
        print(analysis)
        
        # Parse analysis to separate HPI and Assessment/Plan
        hpi = ""
        assessment_plan = ""
        
        # More robust parsing logic
        lines = analysis.split('\n')
        current_section = None
        hpi_lines = []
        assessment_lines = []
        
        for line in lines:
            line_upper = line.upper().strip()
            
            # Detect section headers
            if any(keyword in line_upper for keyword in ['HISTORY OF PRESENT ILLNESS', 'HPI:', '1. HISTORY']):
                current_section = 'hpi'
                hpi_lines.append(line)
            elif any(keyword in line_upper for keyword in ['ASSESSMENT AND PLAN', 'ASSESSMENT & PLAN', '2. ASSESSMENT', 'PLAN:']):
                current_section = 'assessment'
                assessment_lines.append(line)
            elif line_upper.startswith('2.') and 'ASSESSMENT' in line_upper:
                current_section = 'assessment'
                assessment_lines.append(line)
            elif line_upper.startswith('1.') and 'HISTORY' in line_upper:
                current_section = 'hpi'
                hpi_lines.append(line)
            else:
                # Add content to current section
                if current_section == 'hpi' and line.strip():
                    hpi_lines.append(line)
                elif current_section == 'assessment' and line.strip():
                    assessment_lines.append(line)
        
        hpi = '\n'.join(hpi_lines) if hpi_lines else ""
        assessment_plan = '\n'.join(assessment_lines) if assessment_lines else ""
        
        # If parsing failed, try alternative approach
        if not hpi and not assessment_plan:
            # Try to find sections by looking for numbered lists
            parts = analysis.split('\n\n')
            for part in parts:
                part_upper = part.upper()
                if any(keyword in part_upper for keyword in ['HISTORY', 'HPI', 'PRESENT ILLNESS']):
                    hpi = part
                elif any(keyword in part_upper for keyword in ['ASSESSMENT', 'PLAN', 'DIAGNOSIS', 'TREATMENT']):
                    assessment_plan = part
        
        # Find and fill the best template
        filled_template = None
        best_template = find_best_template(hpi, assessment_plan)
        if best_template:
            filled_template = fill_template_with_llm(
                best_template.content,
                hpi,
                assessment_plan
            )
        
        # Save to database
        conversation = MedicalConversation(
            id=conversation_id,
            patient_name=patient_name,
            physician_name=physician_name,
            audio_file_path=audio_path,
            transcription=transcription,
            history_of_present_illness=hpi,
            assessment_plan=assessment_plan
        )
        
        db.session.add(conversation)
        db.session.commit()
        
        return jsonify({
            'conversation_id': conversation_id,
            'transcription': transcription,
            'history_of_present_illness': hpi,
            'assessment_plan': assessment_plan,
            'filled_template': filled_template,
            'message': 'Audio processed successfully'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/upload-document', methods=['POST'])
def upload_document():
    """Upload and process a document for the knowledge base"""
    try:
        if 'document' not in request.files:
            return jsonify({'error': 'No document file provided'}), 400
        
        document_file = request.files['document']
        
        if document_file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Determine file type
        filename = document_file.filename.lower()
        if filename.endswith('.pdf'):
            file_type = 'pdf'
        elif filename.endswith(('.docx', '.doc')):
            file_type = 'docx'
        elif filename.endswith('.txt'):
            file_type = 'txt'
        else:
            return jsonify({'error': 'Unsupported file type. Supported: PDF, DOCX, TXT'}), 400
        
        # Save document
        documents_folder = 'documents'
        if not os.path.exists(documents_folder):
            os.makedirs(documents_folder)
        
        document_path = os.path.join(documents_folder, document_file.filename)
        document_file.save(document_path)
        
        # Process document
        result = document_processor.process_document(document_path, file_type)
        
        if result['success']:
            return jsonify({
                'message': 'Document processed successfully',
                'chunks_created': result['chunks_created'],
                'file_type': result['file_type']
            })
        else:
            return jsonify({'error': result['error']}), 500
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/upload-template', methods=['POST'])
def upload_template():
    """Upload and store an EPIC template (TXT or DOCX) with optional metadata."""
    try:
        if 'template' not in request.files:
            return jsonify({'error': 'No template file provided'}), 400
        template_file = request.files['template']
        name = request.form.get('name', template_file.filename)
        metadata = request.form.get('metadata', '{}')  # JSON string
        # Read file content
        filename = template_file.filename.lower()
        if filename.endswith('.txt'):
            content = template_file.read().decode('utf-8')
        elif filename.endswith('.docx'):
            from docx import Document
            import io
            doc = Document(io.BytesIO(template_file.read()))
            content = '\n'.join([para.text for para in doc.paragraphs])
        else:
            return jsonify({'error': 'Unsupported file type. Supported: TXT, DOCX'}), 400
        # Store in DB
        template_id = str(uuid.uuid4())
        template = Template(
            id=template_id,
            name=name,
            content=content,
            template_metadata=metadata
        )
        db.session.add(template)
        db.session.commit()
        # Add to vector store
        try:
            meta_dict = json.loads(metadata) if metadata else {}
        except Exception:
            meta_dict = {}
        template_vectorstore.add_template(template_id, content, meta_dict)
        return jsonify({'message': 'Template uploaded successfully', 'template_id': template.id})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/knowledge-base/stats', methods=['GET'])
def get_knowledge_base_stats():
    """Get knowledge base statistics"""
    try:
        stats = document_processor.get_knowledge_base_stats()
        return jsonify(stats)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/knowledge-base/clear', methods=['POST'])
def clear_knowledge_base():
    """Clear the knowledge base"""
    try:
        result = document_processor.clear_knowledge_base()
        if result['success']:
            return jsonify({'message': result['message']})
        else:
            return jsonify({'error': result['error']}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/conversations', methods=['GET'])
def get_conversations():
    """Get all conversations"""
    try:
        conversations = MedicalConversation.query.order_by(MedicalConversation.created_at.desc()).all()
        result = []
        
        for conv in conversations:
            result.append({
                'id': conv.id,
                'patient_name': conv.patient_name,
                'physician_name': conv.physician_name,
                'conversation_date': conv.conversation_date.isoformat(),
                'created_at': conv.created_at.isoformat()
            })
        
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/conversations/<conversation_id>', methods=['GET'])
def get_conversation(conversation_id):
    """Get specific conversation details"""
    try:
        conversation = MedicalConversation.query.get(conversation_id)
        if not conversation:
            return jsonify({'error': 'Conversation not found'}), 404
        
        return jsonify({
            'id': conversation.id,
            'patient_name': conversation.patient_name,
            'physician_name': conversation.physician_name,
            'conversation_date': conversation.conversation_date.isoformat(),
            'transcription': conversation.transcription,
            'history_of_present_illness': conversation.history_of_present_illness,
            'assessment_plan': conversation.assessment_plan,
            'created_at': conversation.created_at.isoformat()
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/conversations/<conversation_id>/filled-template', methods=['GET'])
def get_filled_template(conversation_id):
    """Return the best-matched filled template for a conversation."""
    try:
        conversation = MedicalConversation.query.get(conversation_id)
        if not conversation:
            return jsonify({'error': 'Conversation not found'}), 404
        best_template = find_best_template(conversation.history_of_present_illness, conversation.assessment_plan)
        if not best_template:
            return jsonify({'error': 'No suitable template found'}), 404
        filled_template = fill_template_with_llm(
            best_template.content,
            conversation.history_of_present_illness or '',
            conversation.assessment_plan or ''
        )
        return jsonify({'filled_template': filled_template, 'template_id': best_template.id, 'template_name': best_template.name})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def find_best_template(hpi, assessment_plan, top_k=3):
    """Hybrid: Use vector search for top N, then LLM to select the best template."""
    try:
        query_text = (hpi or '') + ' ' + (assessment_plan or '')
        # Step 1: Vector search for top N
        candidates = template_vectorstore.search_templates(query_text, top_k=top_k)
        if not candidates:
            return None
        if len(candidates) == 1:
            # Only one candidate, return it
            template_id = candidates[0].metadata.get('template_id')
            return Template.query.get(template_id)
        # Step 2: LLM selection
        # Prepare prompt
        prompt = f"""
        Given the following conversation summary and template options, select the best template for this case. 
        Conversation summary:
        HPI: {hpi}
        Assessment & Plan: {assessment_plan}

        Template options:
        """
        for i, doc in enumerate(candidates):
            template_id = doc.metadata.get('template_id')
            template = Template.query.get(template_id)
            prompt += f"Option {i+1} (ID: {template_id}, Name: {template.name}):\n{template.content}\n\n"
        prompt += "\nRespond ONLY with the ID of the best template option."
        llm = ChatOpenAI(
            model="gpt-4",
            openai_api_key=os.getenv("OPENAI_API_KEY"),
            temperature=0.0,
            max_tokens=20,
        )
        messages = [
            SystemMessage(content="You are a helpful assistant for template selection."),
            HumanMessage(content=prompt)
        ]
        response = llm.invoke(messages)
        selected_id = response.content.strip()
        # Try to extract a valid template ID from the response
        for doc in candidates:
            template_id = doc.metadata.get('template_id')
            if template_id in selected_id:
                return Template.query.get(template_id)
        # Fallback: return top candidate
        template_id = candidates[0].metadata.get('template_id')
        return Template.query.get(template_id)
    except Exception as e:
        print(f"Error in hybrid template matching: {e}")
        return None

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5100)
